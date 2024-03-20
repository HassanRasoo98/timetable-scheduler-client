# all the methods/functions used in the app.py file have been defined in this file

import os
import re
import pytz
import shutil
import requests
import pandas as pd
import streamlit as st
from docx import Document
from datetime import datetime, timedelta, timezone
from streamlit_star_rating import st_star_rating

def get_base_url():
    # return "http://127.0.0.1:5000/" # development url
    return "https://hassanrasool.pythonanywhere.com/" # production url

# Function to update the timetable
def get_last_update_time(base_url):
    # st.info("Updating timetable...")  # You can customize the message
    update_url = base_url + "get_modification_time"
    response = requests.get(update_url)
    # print(response.json())
    if response.status_code == 200:
        st.success("Timetable last updated on {}".format(response.json()))
    else:
        st.error(f"Failed to fetch timetable last update date. Status code: {response.status_code}")

    return allow_update(response.json())

def results_handler(result, folder='results'):
    print(result)
    # Create a subfolder named 'results' in the current working directory
    results_folder = folder
    
    # remove previouisly fetched results to make space for new ones
    if os.path.exists(os.path.join(os.getcwd(), results_folder)):
        shutil.rmtree(results_folder)
        # print('file removed')
        
    os.makedirs(results_folder, exist_ok=True)
    
    result.drop_duplicates(inplace=True, ignore_index=True)
    
    try:
        result['Start_Time'] = result['Start_Time'].apply(format_time)
        # result['End_Time'] = result['End_Time'].apply(format_time)
        result['Start_Time'] = pd.to_datetime(result['Start_Time'], format='%I:%M %p')

        # Group by 'Day', sort each group, and save as a separate CSV file in the 'results' subfolder
        for day, group in result.groupby('Day'):
            sorted_group = group.sort_values(by='Start_Time')
            csv_file_path = os.path.join(results_folder, f'{day.replace(".xlsx", "")}.csv')
            sorted_group.drop(['Day', 'Start_Time', 'End_Time'], axis=1, inplace=True)
            
            # Desired column order
            desired_order = ['Subject', 'Class', 'Time']

            # Rearrange columns
            sorted_group = sorted_group[desired_order]
            sorted_group.to_csv(csv_file_path, index=False)

        # print("CSV files saved successfully in the 'results' subfolder after sorting and converting time to datetime.")
    
    except:
        '''
            exam case
        '''
        for date, group in result.groupby('Days & Date'):
            csv_file_path = os.path.join(results_folder, f'{date}.csv')
            group.drop(['Days & Date'], axis=1, inplace=True)
            
            desired_order = ['Subject', 'Day_Name', 'Time']
            # Rearrange columns
            sorted_group = group[desired_order]
            sorted_group.to_csv(csv_file_path, index=False)
            
def get_draft_version():
    version_url = get_base_url() + 'get_version'
    try:
        response = requests.get(version_url)
        response.raise_for_status()  # Raise an exception for 4XX or 5XX status codes
        data = response.json()
        return data.get('filename')
    
    except requests.exceptions.RequestException as e:
        print(f"Error fetching draft version: {e}")
        # You might want to handle or log the error in a better way
        return None
            
# Function to update the timetable
def update_timetable(base_url):
    st.info("Updating timetable...")  # You can customize the message
    update_url = base_url + "update-timetable"
    response = requests.get(update_url)
    if response.status_code == 200:
        st.success("Timetable updated successfully.")
    else:
        st.error(f"Failed to update timetable. Fatching timetable from previous version. Status code: {response.status_code}")

# Function to read CSV files from a folder
def read_csv_files(folder_path, sort=False):
    csv_files = [file for file in os.listdir(folder_path) if file.endswith('.csv')]
    
    # sort the files according to the day names
    if sort:
        csv_files = sorted(csv_files, key=order_files)
    else:
        csv_files = sorted(csv_files)        
    
    # print('csv files : ', csv_files)
    dfs = {file: pd.read_csv(os.path.join(folder_path, file)) for file in csv_files}
    
    return dfs

# Initialize session_state
def initialize_session_state():
    if 'selected_tab' not in st.session_state:
        st.session_state.selected_tab = None
        
def order_files(file):
    order = {
        "Monday.csv": 1,
        "Tuesday.csv": 2,
        "Wednesday.csv": 3,
        "Thursday.csv": 4,
        "Friday.csv": 5,
        "Saturday.csv": 6,
        "Sunday.csv": 7
    }
    return order.get(file, 0)
  
def extract_numbers(input_string):
    # Use regex to find all numeric values in the string
    numbers = re.findall(r'\d+', input_string)
    return [int(number) for number in numbers]

def format_time(time):
    hour = extract_numbers(time)[0]

    if hour >= 8 and hour < 12:
        return time + ' AM'
        
    else:
        return time + ' PM'
         
def allow_update(provided_datetime_str):
    # print(provided_datetime_str)
    provided_datetime = datetime.strptime(provided_datetime_str, "%a, %d %b %Y %H:%M:%S %Z")
    provided_datetime = datetime.strptime(provided_datetime_str, "%a, %d %b %Y %H:%M:%S %Z")
    provided_datetime = provided_datetime.replace(tzinfo=timezone.utc)

    # Current datetime with timezone
    current_datetime = datetime.now(timezone.utc)
    # print(current_datetime)

    # Calculate the difference
    time_difference = current_datetime - provided_datetime
    # Convert to total seconds
    total_seconds = time_difference.total_seconds()

    # Format the time in a human-readable way
    hours, remainder = divmod(total_seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    minutes = int(minutes)
    seconds = int(seconds)

    # print("Time difference:", time_difference)
    
    time_string = "{} m {} s".format(minutes, seconds)

    # Specify the maximum allowed duration (30 minutes)
    max_allowed_duration = timedelta(minutes=30)

    # Check if the difference is less than half an hour
    if time_difference < max_allowed_duration:
        return False, time_string # do not allow update, last update was less than 30 minutes ago
    else:
        return True, time_string

def create_file(folder='results'):
    # Specify the subfolder containing CSV files
    csv_folder = folder

    # Create a Word document
    doc = Document()
    doc.add_heading('Timetable Spring 2024', level=0)  # Add the title
    
    files = os.listdir(csv_folder)
    files = sorted(files, key=order_files)

    # Loop through all CSV files in the subfolder
    for csv_file in files:
        # Read the CSV file into a Pandas DataFrame
        df = pd.read_csv(os.path.join(csv_folder, csv_file))

        title = csv_file.split('.')[0]
        # Add a section for the CSV file name
        doc.add_heading(title, level=1)

        # Add a table to the Word document
        num_rows, num_cols = df.shape
        table = doc.add_table(rows=num_rows + 1, cols=num_cols)
        table.autofit = True

        # Add column headers to the table
        for col_num, col_name in enumerate(df.columns):
            table.cell(0, col_num).text = col_name

        # Add data to the table
        for row_num, row_data in enumerate(df.itertuples(), start=1):
            for col_num, value in enumerate(row_data[1:], start=0):
                table.cell(row_num, col_num).text = str(value)
                
    # Add the concluding lines
    doc.add_paragraph("\nThis Timetable was made automatically using the tool available at: https://isb-fastnuces-timetable-scheduler.streamlit.app/\nIf you enjoyed using this app, please provide feedback. "
                      "If you want to support upcoming projects like this one, "
                      "you can motivate me financially :)\n")
    doc.add_paragraph("01397991921003 - HBL")
                
    if os.path.exists('combined_data.docx'):
        os.remove('combined_data.docx')

    # if os.path.exists('combined_data.pdf'):
    #     os.remove('combined_data.pdf')

    # Save the Word document
    doc.save('combined_data.docx')
    # print('Word document saved.')

    # # Convert Word document to PDF
    # convert("combined_data.docx", "combined_data.pdf")
    # print('PDF document saved.')
    
# Function to fetch current rating from the server
def fetch_current_rating(base_url):
    # Assuming your server provides the current rating at the '/current-rating' endpoint
    current_rating_url = base_url + "/current-rating"
    response = requests.get(current_rating_url)
    if response.status_code == 200:
        rating = response.json()["rating"]
        rating = round(rating, 2)
        return rating
    else:
        st.error("Failed to fetch current rating.")
        return 0
    

def send_rating_to_server(rating, base_url):
    # Assuming you have a server endpoint to send the rating via POST request
    endpoint = base_url + 'submit-rating'
    data = {"rating": rating}
    response = requests.post(endpoint, json=data)
    # if response.status_code == 200:
        # st.success('Thank you for rating us.')
        
# Display the fetched rating on the top left corner using st.sidebar
def show_rating(base_url):
    # Fetch the current rating from the server
    current_rating, total = fetch_current_rating(base_url)
    with st.sidebar:
        st.write("Current Rating:")
        st.write(current_rating)
        
        st.write("Total Votes")
        st.write(total)


def get_user_rating_handler(base_url):
    st.divider()
    stars = st_star_rating("Please rate your experience", size = 20, maxValue=5, defaultValue=None, key="rating")
    st.divider()
    send_rating_to_server(stars, base_url)
    
def subscribe_to_updates(base_url, email):
    email_url = base_url + 'subscribe-email'
    data = {"email": email}
    response = requests.post(email_url, json=data)
    if response.status_code == 200:
        st.success('Thank you for subscribing. You will now get notified by email about new updates!')

def store_feedback(feedback_type, rating, feedback_text, contact_email, 
                   name, batch, department, roll_num):
    
    # For demonstration, you can store feedback in a CSV file
    feedback_data = {'Feedback Type': [feedback_type],
                     'Feedback': [feedback_text],
                     'Contact Email': [contact_email],
                     'Name': name,
                     'Batch': batch,
                     'Department': department,
                     'Roll Number': roll_num,
                     'Rating': rating}
    
    send_rating_to_server(rating, get_base_url())
        
    # send to server to store the file
    feedback_url = get_base_url() + 'post-feedback'
    response = requests.post(feedback_url, json=feedback_data)
    
    if response.status_code == 200:
        st.success('Your feedback has been recorded.')
    else:
        st.error(f'Error in submitting feedback. Status {response.status_code}')
    
def email_checker(email):
    # Regular expression pattern for validating email addresses
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    
    # Check if the email matches the pattern
    if re.match(pattern, email):
        return True
    else:
        return False

def tab1_empty_classroom_handler():
    base_url = get_base_url()
    # Fetch a list of all files
    all_files_url = base_url + "get_files"
    response = requests.get(all_files_url)
    files = None

    if response.status_code == 200:
        files = response.json()
    else:
        st.error(f"Failed to fetch files. Status code: {response.status_code}")
        st.stop()
        
        # Create a multi-select widget for selecting subjects
    selected_file = st.selectbox("Select Day:", files)   
    # print('selected file : ', selected_file)

    # Create two radio buttons side by side
    stype = st.radio('Find in : ', ('Room', 'Lab'))
    url = base_url + "selected-file"
    payload = {"file": selected_file, "selection_type": stype}
    response1 = requests.post(url, json=payload)    
    timeslots = response1.json()
    # print(timeslots)

    with st.form(key="form"):
        selected_timeslot = st.selectbox("Select a timeslot:", timeslots)
        submit_button = st.form_submit_button(label="Find Empty Rooms")
        # print(f'selected_timeslot {selected_timeslot}')  
        
    if submit_button:
        # print(f'selected_timeslot after clicking button {selected_timeslot}') 
        # print('button clicked')
        free_room_url = base_url + "get-free-room"
        # print(free_room_url)
        payload = {"timeslot": selected_timeslot}
        free_room_url_response = requests.post(free_room_url, json=payload)
        
        if free_room_url_response.status_code == 200:
            result = free_room_url_response.json()
            df = pd.DataFrame(result, columns=['Free Rooms'])

            st.table(df)
            
        else:
            st.error(f"Failed to fetch timetable. Status code: {free_room_url_response.status_code}")
            # print(free_room_url_response)
            
    st.divider()


def tab2_empty_classroom_handler():
    # Get current date and time
    current_date_time = datetime.now()
    
    # Define time zone for Pakistan
    pakistan_time_zone = pytz.timezone('Asia/Karachi')
    
    # Convert current date and time to Pakistan time
    current_date_time_pakistan = current_date_time.astimezone(pakistan_time_zone)
    
    # Extract current day and time in Pakistan time
    current_day = current_date_time_pakistan.strftime("%A")
    time = current_date_time_pakistan.strftime('%I:%M%p')
    
    print("Current day in Pakistan:", current_day)
    print("Current time in Pakistan:", time)
    
    # # Get current date and time
    # current_date_time = datetime.now()

    # # Extract current day and time
    # current_day = current_date_time.strftime("%A")
    # time = current_date_time.strftime('%I:%M%p')

    # print("Current day:", current_day)
    # print("Current time:", time)
    
    now = {'current-day': current_day.strip(), 'current-time': time}
    
    # send this data to API endpoint
    now_url = get_base_url() + 'now-empty'
    response = requests.post(now_url, json=now)
    
    if response.status_code == 200:
        classes = response.json()['result1']
        labs = response.json()['result2']
        
        st.subheader('Free Classes')
        # Display upcoming features
        for class_ in classes:
            st.write(f"- {class_}")
        
        st.subheader('Free Labs')
        # Display upcoming features
        for lab in labs:
            st.write(f"- {lab}")
            
    # current time does not match any time slots (against office timings)
    elif response.status_code == 201:
        st.error(response.json()['message'])
        
    # any other error
    else:
        st.error('An error occured')
        print(response.text)
